from docx import Document
from docx.shared import Pt
import re

def replace_about_me(docx_path, new_summary, output_path="updated_cv.docx"):
    doc = Document(docx_path)
    
    summary_headers = ["professional summary", "about me", "summary", "profile", "objective"]
    
    found_summary = False
    
    for i, paragraph in enumerate(doc.paragraphs):
        text_lower = paragraph.text.lower().strip()
        
        if any(header in text_lower for header in summary_headers):
            found_summary = True
            
            if i + 1 < len(doc.paragraphs):
                doc.paragraphs[i + 1].text = new_summary
                
                summary_paragraphs = new_summary.split('\n\n')
                
                if len(summary_paragraphs) > 1:
                    current_para = doc.paragraphs[i + 1]
                    current_para.text = summary_paragraphs[0]
                    
                    # Insert additional paragraphs after the first one
                    for j, para_text in enumerate(summary_paragraphs[1:], 1):
                        if para_text.strip():
                            new_para = doc.add_paragraph(para_text.strip())
                            # Move paragraph to correct position in document structure
                            doc._body._body.insert(current_para._element.getnext(), new_para._element)
                            current_para = new_para
                else:
                    doc.paragraphs[i + 1].text = new_summary
                
                break
    
    if not found_summary:
        new_para = doc.add_paragraph("Professional Summary")
        new_para.style = 'Heading 2'
        
        summary_para = doc.add_paragraph(new_summary)
        
        doc._body._body.insert(0, new_para._element)
        doc._body._body.insert(1, summary_para._element)
    
    doc.save(output_path)
    print(f"✅ CV saved as {output_path}")

def replace_summary_section_docx(docx_path, new_summary, output_path="updated_cv.docx"):
    doc = Document(docx_path)
    
    summary_headers = ["professional summary", "about me", "summary", "profile", "objective"]
    
    summary_start = None
    
    for i, paragraph in enumerate(doc.paragraphs):
        text_lower = paragraph.text.lower().strip()
        
        if any(header in text_lower for header in summary_headers):
            summary_start = i
            break
    
    if summary_start is not None and summary_start + 1 < len(doc.paragraphs):
        # Clear existing content and set new text with Times New Roman 12
        para = doc.paragraphs[summary_start + 1]
        para.clear()
        run = para.add_run(new_summary)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        print(f"✅ Replaced summary at paragraph {summary_start + 1}")
    else:
        new_doc = Document()
        
        header_para = new_doc.add_paragraph("Professional Summary")
        header_para.style = 'Heading 2'
        
        summary_para = new_doc.add_paragraph()
        run = summary_para.add_run(new_summary)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        for paragraph in doc.paragraphs:
            new_doc.add_paragraph(paragraph.text)
        
        new_doc.save(output_path)
        print(f"✅ Created new CV with summary at the top")
        return
    
    doc.save(output_path)
    print(f"✅ CV saved as {output_path}")
